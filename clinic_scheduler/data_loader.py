"""Data loading and clinic configuration helpers."""

from __future__ import annotations

import logging
import math
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from docx import Document

LOGGER = logging.getLogger(__name__)

SLOT_MINUTES = 5
ROOMS = list(range(1, 17))
DAY_NAMES = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
DAY_TO_INDEX = {day: idx for idx, day in enumerate(DAY_NAMES)}
INDEX_TO_DAY = {idx: day for day, idx in DAY_TO_INDEX.items()}

OPERATING_HOURS = {
    0: (540, 1020),
    1: (540, 1020),
    2: (510, 1110),
    3: (510, 990),
    4: (510, 940),
}

BLOCKED_PERIODS = {
    0: [(540, 570), (690, 720), (720, 780), (990, 1020)],
    1: [(540, 570), (690, 720), (720, 780), (990, 1020)],
    2: [(540, 570), (690, 720), (720, 780), (990, 1020)],
    3: [(540, 570), (690, 720), (720, 780), (990, 1020)],
    4: [(480, 510), (690, 720), (720, 780), (900, 930)],
}


def _normalize_string(value: Any) -> str:
    """Return a stripped string representation."""
    if pd.isna(value):
        return ""
    return str(value).strip()


def time_to_minutes(time_str: str) -> int:
    """Convert HH:MM:SS string to minutes from midnight."""
    parts = [int(part) for part in time_str.split(":")]
    hours, minutes = parts[:2]
    return hours * 60 + minutes


def floor_to_slot(minute_value: int) -> int:
    """Round a minute value down to the nearest slot boundary."""
    return (minute_value // SLOT_MINUTES) * SLOT_MINUTES


def ceil_to_slot(minute_value: int) -> int:
    """Round a minute value up to the nearest slot boundary."""
    return int(math.ceil(minute_value / SLOT_MINUTES) * SLOT_MINUTES)


def get_operating_hours(day_of_week: int) -> tuple[int, int]:
    """Return operating hours for a weekday index."""
    return OPERATING_HOURS[day_of_week]


def get_blocked_periods(day_of_week: int, policy_params: dict[str, Any] | None = None) -> list[tuple[int, int]]:
    """Return blocked periods for a weekday, optionally modified by policy."""
    if not policy_params:
        return []

    if not (policy_params.get("respect_blocked_periods") or policy_params.get("allow_admin_overflow")):
        return []

    blocked = list(BLOCKED_PERIODS[day_of_week])

    if policy_params.get("allow_admin_overflow"):
        morning_admin = (540, 570) if day_of_week < 4 else (480, 510)
        allowed_admin_minutes = {
            floor_to_slot(int(minute))
            for minute in policy_params.get("admin_overflow_minutes", set())
        }
        if allowed_admin_minutes:
            new_blocked: list[tuple[int, int]] = []
            for start, end in blocked:
                if (start, end) != morning_admin:
                    new_blocked.append((start, end))
                    continue
                segments: list[tuple[int, int]] = []
                cursor = start
                allowed_sorted = sorted(slot for slot in allowed_admin_minutes if start <= slot < end)
                for slot in allowed_sorted:
                    if cursor < slot:
                        segments.append((cursor, slot))
                    cursor = min(end, slot + SLOT_MINUTES)
                if cursor < end:
                    segments.append((cursor, end))
                new_blocked.extend(segment for segment in segments if segment[0] < segment[1])
            blocked = new_blocked
    return blocked


def is_time_window_feasible(
    start_min: int,
    slots: int,
    day_of_week: int,
    policy_params: dict[str, Any] | None = None,
) -> bool:
    """Check whether an appointment fits within open hours and avoids blocked periods."""
    open_min, close_min = get_operating_hours(day_of_week)
    end_min = start_min + slots * SLOT_MINUTES
    if start_min < open_min or end_min > close_min:
        return False

    blocked_periods = get_blocked_periods(day_of_week, policy_params)
    for slot_start in range(start_min, end_min, SLOT_MINUTES):
        slot_end = slot_start + SLOT_MINUTES
        for blocked_start, blocked_end in blocked_periods:
            if slot_start < blocked_end and slot_end > blocked_start:
                return False
    return True


def iter_feasible_starts(
    day_of_week: int,
    slots: int,
    policy_params: dict[str, Any] | None = None,
) -> list[int]:
    """Enumerate feasible slot-aligned start minutes for the given duration.

    Deprecated for the fixed-time room-assignment model; retained for compatibility.
    """
    open_min, close_min = get_operating_hours(day_of_week)
    feasible: list[int] = []
    for start_min in range(open_min, close_min, SLOT_MINUTES):
        if is_time_window_feasible(start_min, slots, day_of_week, policy_params):
            feasible.append(start_min)
    return feasible


def load_appointments(csv_path: str) -> pd.DataFrame:
    """
    Load and clean appointment CSV.

    Returns a DataFrame with standardized appointment fields.
    """
    raw_df = pd.read_csv(csv_path)
    raw_df.columns = [str(column).strip() for column in raw_df.columns]
    object_columns = raw_df.select_dtypes(include=["object", "string"]).columns
    for column in object_columns:
        raw_df[column] = raw_df[column].map(_normalize_string)

    raw_df = raw_df[
        (raw_df["Cancelled Appts"] != "Y") & (raw_df["Deleted Appts"] != "Y")
    ].copy()
    raw_df["Appt Date"] = pd.to_datetime(raw_df["Appt Date"], format="%m-%d-%Y")
    raw_df["start_min"] = raw_df["Appt Time"].map(time_to_minutes)
    raw_df["no_show"] = raw_df["No Show Appts"].eq("Y")
    raw_df["slots"] = raw_df["Appt Duration"].astype(int).clip(lower=0).map(
        lambda duration: max(1, math.ceil(duration / SLOT_MINUTES))
    )

    cleaned_df = raw_df.rename(
        columns={
            "Patient Id": "patient_id",
            "Appt Date": "date",
            "Primary Provider": "provider",
            "ApptStatusSingleView": "status",
            "Appt Duration": "duration",
            "Appt Type": "appt_type",
        }
    )[
        [
            "patient_id",
            "date",
            "provider",
            "status",
            "no_show",
            "start_min",
            "duration",
            "appt_type",
            "slots",
        ]
    ].copy()
    cleaned_df["day_of_week"] = cleaned_df["date"].dt.dayofweek
    cleaned_df = cleaned_df[cleaned_df["duration"].astype(int) > 0].copy()
    cleaned_df["appt_type"] = cleaned_df["appt_type"].fillna("").astype(str).str.strip()
    cleaned_df["status"] = cleaned_df["status"].fillna("").astype(str).str.strip()
    cleaned_df["patient_id"] = cleaned_df["patient_id"].astype(str).str.strip()
    cleaned_df["provider"] = cleaned_df["provider"].astype(str).str.strip()
    cleaned_df["date_str"] = cleaned_df["date"].dt.strftime("%Y-%m-%d")
    cleaned_df["appointment_id"] = cleaned_df.apply(
        lambda row: (
            f"{row['date_str']}|{row['provider']}|{row['patient_id']}|"
            f"{int(row['start_min'])}|{int(row['duration'])}|{row.name}"
        ),
        axis=1,
    )
    return cleaned_df.reset_index(drop=True)


def _parse_room_token(room_text: str) -> int | None:
    """Parse a room id from a room token."""
    match = re.search(r"(?:ROOM|RM)\s*(\d+)", room_text.upper())
    if match:
        return int(match.group(1))
    return None


def _parse_assignment_cell(cell_text: str) -> dict[str, Any]:
    """Parse a single room assignment cell into AM/PM availability."""
    normalized = _normalize_string(cell_text).upper()
    if not normalized or normalized == "N/A":
        return {"am_room": None, "pm_room": None, "available": False, "status": "N/A"}
    if normalized == "CLOSED":
        return {"am_room": None, "pm_room": None, "available": False, "status": "CLOSED"}
    if normalized == "NO ROOM AVAILABLE":
        return {"am_room": None, "pm_room": None, "available": True, "status": "NO ROOM AVAILABLE"}

    if "/" in normalized:
        am_room = None
        pm_room = None
        for part in normalized.split("/"):
            room_id = _parse_room_token(part)
            if room_id is None:
                continue
            if "(AM)" in part:
                am_room = room_id
            elif "(PM)" in part:
                pm_room = room_id
            else:
                am_room = room_id
                pm_room = room_id
        return {"am_room": am_room, "pm_room": pm_room, "available": True, "status": "SPLIT"}

    room_id = _parse_room_token(normalized)
    if room_id is None:
        LOGGER.warning("Unrecognized room assignment cell: %s", cell_text)
        return {"am_room": None, "pm_room": None, "available": False, "status": "UNKNOWN"}
    if "(AM)" in normalized:
        return {"am_room": room_id, "pm_room": None, "available": True, "status": "AM_ONLY"}
    if "(PM)" in normalized:
        return {"am_room": None, "pm_room": room_id, "available": True, "status": "PM_ONLY"}
    return {"am_room": room_id, "pm_room": room_id, "available": True, "status": "FULL_DAY"}


def load_room_assignments(docx_path: str) -> dict[str, dict[str, dict[str, Any]]]:
    """
    Parse provider-room assignment data from the DOCX table.

    Returns nested mapping: assignments[provider][day_name] -> room assignment metadata.
    """
    document = Document(docx_path)
    if not document.tables:
        raise ValueError(f"No tables found in DOCX file: {docx_path}")

    table = document.tables[0]
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    header = rows[0]
    day_columns = {column: index for index, column in enumerate(header) if column in DAY_NAMES}
    assignments: dict[str, dict[str, dict[str, Any]]] = {}

    for row in rows[1:]:
        provider = _normalize_string(row[0])
        if not provider:
            continue
        assignments[provider] = {}
        for day_name in DAY_NAMES:
            cell_idx = day_columns.get(day_name)
            cell_text = row[cell_idx] if cell_idx is not None and cell_idx < len(row) else ""
            assignments[provider][day_name] = _parse_assignment_cell(cell_text)
    return assignments


def build_distance_matrix() -> np.ndarray:
    """Return the 16x16 symmetric room distance matrix."""
    upper_triangle = [
        [0, 0.5, 1.5, 2.5, 4, 4.7, 2, 3, 4.7, 4.1, 3.3, 1.5, 3, 6.1, 11.5, 12],
        [0, 1, 2, 3.5, 4.2, 1.5, 2.5, 4.2, 3.6, 2.8, 2, 3.5, 5.6, 11, 11.5],
        [0, 1, 2.5, 3.2, 0.5, 1.5, 3.2, 2.6, 1.8, 3, 4.5, 4.6, 10, 10.5],
        [0, 1.5, 2.2, 1.5, 1.5, 2.2, 1.6, 2.8, 4, 5.5, 3.6, 9, 9.5],
        [0, 0.7, 3, 3, 2.5, 3.1, 4.3, 5.5, 6, 3.9, 7.5, 8],
        [0, 3.7, 3.7, 3.2, 3.8, 5, 6.2, 6.7, 4.6, 6.8, 7.3],
        [0, 1, 2.7, 2.1, 1.3, 2.5, 4, 4.1, 9.5, 10],
        [0, 1.7, 1.1, 1.3, 2.5, 4, 3.1, 8.5, 9],
        [0, 0.6, 1.8, 3.2, 5.7, 1.4, 6.8, 7.3],
        [0, 1.2, 2.6, 5.1, 2, 7.4, 7.9],
        [0, 1.8, 4.3, 2.8, 8.2, 8.7],
        [0, 2.5, 4.6, 10, 10.5],
        [0, 7.1, 12.5, 13],
        [0, 5.4, 5.9],
        [0, 0.5],
        [0],
    ]
    matrix = np.zeros((16, 16), dtype=float)
    for row_idx, row in enumerate(upper_triangle):
        for offset, value in enumerate(row):
            col_idx = row_idx + offset
            matrix[row_idx, col_idx] = value
            matrix[col_idx, row_idx] = value
    return matrix


@dataclass(frozen=True)
class ProviderDayKey:
    """A provider-day-date key."""

    provider: str
    day: str
    date_str: str


def assigned_room_for_time(
    room_assignments: dict[str, dict[str, dict[str, Any]]],
    provider: str,
    day_name: str,
    start_min: int,
) -> int | None:
    """Return the assigned room for a provider/day/time, if one exists."""
    assignment = room_assignments.get(provider, {}).get(day_name, {})
    if start_min < 720:
        return assignment.get("am_room") or assignment.get("pm_room")
    return assignment.get("pm_room") or assignment.get("am_room")


def build_historical_schedule(
    appointments_df: pd.DataFrame,
    room_assignments: dict[str, dict[str, dict[str, Any]]],
) -> pd.DataFrame:
    """
    Build a baseline schedule dataframe directly from historical appointments.

    The resulting schema matches the solver's ``schedule_df`` so the same
    visualization code can render historical and optimized schedules.
    """
    if appointments_df.empty:
        return pd.DataFrame(
            columns=[
                "provider",
                "day",
                "date",
                "appointment_id",
                "patient_id",
                "start_min",
                "end_min",
                "occupied_end_min",
                "room",
                "slots",
                "buffered_slots",
                "is_phantom",
            ]
        )

    rows: list[dict[str, Any]] = []
    sorted_df = appointments_df.sort_values(["date", "provider", "start_min", "patient_id"]).copy()
    for _, row in sorted_df.iterrows():
        day_name = row["date"].day_name()
        start_min = int(row["start_min"])
        slots = int(row["slots"])
        room = assigned_room_for_time(room_assignments, str(row["provider"]), day_name, start_min)
        end_min = start_min + slots * SLOT_MINUTES
        rows.append(
            {
                "provider": str(row["provider"]),
                "day": day_name,
                "date": row["date"].strftime("%Y-%m-%d"),
                "appointment_id": str(row["appointment_id"]),
                "patient_id": str(row["patient_id"]),
                "start_min": start_min,
                "end_min": end_min,
                "occupied_end_min": end_min,
                "room": room,
                "slots": slots,
                "buffered_slots": slots,
                "is_phantom": bool(row.get("is_phantom", False)),
            }
        )

    schedule_df = pd.DataFrame(rows)
    return schedule_df.sort_values(["date", "provider", "start_min", "patient_id"]).reset_index(drop=True)
